import numpy as np
import pandas as pd
import math
import scipy.stats as stats
from docx import Document
import matplotlib.pyplot as plt


N = 10
P = 0.351
SAMPLE_SIZE = 200
ALPHA = 0.05
SEED = 1011

pd.set_option('display.float_format', '{:.5f}'.format)
np.set_printoptions(precision=5, suppress=True)


def binomial_probabilities(n: int, p: float) -> list[float]:
    return [math.comb(n, k) * p**k * (1 - p) ** (n - k) for k in range(n + 1)]


def standard_method(probabilities: list[float], size: int) -> np.ndarray:
    np.random.seed(SEED)
    return np.searchsorted(np.cumsum(probabilities), np.random.random(size))


def scipy_method(n: int, p: float, size: int) -> np.ndarray:
    return stats.binom.rvs(n, p, size=size, random_state=SEED + 1)


def build_stat_table(
    sample: np.ndarray, probabilities: list[float], n: int, size: int
) -> pd.DataFrame:
    counts = np.bincount(sample, minlength=n + 1)

    return pd.DataFrame(
        {
            'x_i': np.arange(n + 1),
            'n_i': counts,
            'w_i': counts / size,
            'p_i': probabilities,
            's_i': np.cumsum(probabilities),
        }
    )


def chi_square_table(stat_table: pd.DataFrame, size: int) -> pd.DataFrame:
    data = stat_table[['w_i', 'p_i']].copy()

    data['|w_i - p_i|'] = np.abs(data['w_i'] - data['p_i'])
    data['chi criteria'] = size * (data['w_i'] - data['p_i']) ** 2 / data['p_i']

    return data


def chi_square_test(chi_table: pd.DataFrame, alpha: float):

    chi_sum = chi_table['chi criteria'].sum()
    chi_crit = stats.chi2.ppf(1 - alpha, len(chi_table) - 1)
    conclusion = (
        'Гипотеза принимается' if chi_sum < chi_crit else 'Гипотеза отвергается'
    )

    print(f'наблюдаемое = {chi_sum:.5f}')
    print(f'критическое = {chi_crit:.5f}')
    print(f'Вывод: {conclusion}')


def uniformity_table(table_1: pd.DataFrame, table_2: pd.DataFrame) -> pd.DataFrame:
    data = pd.DataFrame()

    data['w_i1'] = table_1['w_i']
    data['w_i2'] = table_2['w_i']

    data['uniformity'] = np.where(
        (data['w_i1'] + data['w_i2']) > 0,
        (data['w_i1'] ** 2 + data['w_i2'] ** 2) / (data['w_i1'] + data['w_i2']),
        0.0,
    )

    return data


def uniformity_test(uni_table: pd.DataFrame, size: int, alpha: float):
    column = uni_table['uniformity']
    criteria = 2 * size * (column.sum() - 1)
    chi_crit = stats.chi2.ppf(1 - alpha, len(column) - 1)
    conclusion = (
        'Гипотеза принимается' if criteria < chi_crit else 'Гипотеза отвергается'
    )

    print(f'наблюдаемое = {criteria:.5f}')
    print(f'критическое = {chi_crit:.5f}')
    print(f'Вывод: {conclusion}')


def statistic_analys_table(
    sample: np.ndarray, mean_theoretical: float, var_theoretical: float
) -> pd.DataFrame:
    mean_exp = np.mean(sample)
    var_exp = np.var(sample)
    abs_mean = abs(mean_exp - mean_theoretical)
    abs_var = abs(var_exp - var_theoretical)

    return pd.DataFrame(
        {
            'Название показателя': ['Выборочное среднее', 'Выборочная дисперсия'],
            'Экспериментальное значение': [mean_exp, var_exp],
            'Теоретическое значение': [mean_theoretical, var_theoretical],
            'Абсолютное отклонение': [abs_mean, abs_var],
            'Относительное отклонение': [
                abs_mean / abs(mean_theoretical),
                abs_var / abs(var_theoretical),
            ],
        }
    )


def reshape_20x10(sample: np.ndarray) -> pd.DataFrame:
    return pd.DataFrame(sample.reshape(20, 10))


def plot_polygons_matplotlib(
    table_std: pd.DataFrame, table_scipy: pd.DataFrame, save_path: str = None
):
    x = table_std['x_i']
    w_std = table_std['w_i']
    w_scipy = table_scipy['w_i']
    p = table_std['p_i']

    plt.figure(figsize=(8, 5))

    plt.plot(x, w_std, marker='o', color='blue', linewidth=2, label='Выборка 1')

    plt.plot(x, w_scipy, marker='o', color='green', linewidth=2, label='Выборка 2')

    plt.plot(
        x, p, marker='o', color='red', linewidth=2, label='Теоретические вероятности'
    )

    plt.xlabel('x_i')
    plt.ylabel('Значение')
    plt.xticks(x)
    plt.grid(True, linestyle='--', alpha=0.7)
    plt.legend()
    plt.tight_layout()

    if save_path:
        plt.savefig(save_path, dpi=300)
    plt.show()


def add_table_to_doc(doc: Document, df: pd.DataFrame, title: str):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])

    for j, col in enumerate(df.columns):
        table.rows[0].cells[j].text = str(col)

    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            value = df.iloc[i, j]
            if pd.isna(value):
                table.rows[i + 1].cells[j].text = ''
            elif isinstance(value, (float, np.floating)):
                table.rows[i + 1].cells[j].text = f'{float(value):.5f}'
            else:
                table.rows[i + 1].cells[j].text = str(value)


probabilities = binomial_probabilities(N, P)

sample_std = standard_method(probabilities, SAMPLE_SIZE)
sample_scipy = scipy_method(N, P, SAMPLE_SIZE)

table_std = build_stat_table(sample_std, probabilities, N, SAMPLE_SIZE)
table_scipy = build_stat_table(sample_scipy, probabilities, N, SAMPLE_SIZE)

chi_std_table = chi_square_table(table_std, SAMPLE_SIZE)
chi_scipy_table = chi_square_table(table_scipy, SAMPLE_SIZE)

uniform_table = uniformity_table(table_std, table_scipy)

stats_std = statistic_analys_table(sample_std, N * P, N * P * (1 - P))
stats_scipy = statistic_analys_table(sample_scipy, N * P, N * P * (1 - P))

# Output
print('=== Стандартный метод ===')
print(np.sort(sample_std).reshape(20, 10))
print(table_std)
print(chi_std_table)
print(chi_std_table.sum())
print(stats_std)
chi_square_test(chi_std_table, ALPHA)

print('\n=== Scipy метод ===')
print(np.sort(sample_scipy).reshape(20, 10))
print(table_scipy)
print(chi_scipy_table)
print(chi_scipy_table.sum())
print(stats_scipy)
chi_square_test(chi_scipy_table, ALPHA)

print('\n=== Однородность ===')
print(uniform_table)
uniformity_test(uniform_table, SAMPLE_SIZE, ALPHA)


plot_polygons_matplotlib(table_std, table_scipy, save_path='polygon_distribution.png')

doc = Document()
add_table_to_doc(
    doc, pd.DataFrame(sample_std.reshape(20, 10)), 'Выборка (стандартный метод)'
)
add_table_to_doc(
    doc,
    pd.DataFrame(np.sort(sample_std).reshape(20, 10)),
    'Отсортированная выборка (стандартный метод)',
)

add_table_to_doc(doc, table_std, 'Статистический ряд (стандартный метод)')
add_table_to_doc(doc, chi_std_table, 'Критерий Пирсона (стандартный метод)')
add_table_to_doc(doc, stats_std, 'Числовые характеристики (стандартный метод)')

add_table_to_doc(doc, pd.DataFrame(sample_scipy.reshape(20, 10)), 'Выборка (scipy)')
add_table_to_doc(
    doc,
    pd.DataFrame(np.sort(sample_scipy).reshape(20, 10)),
    'Отсортированная выборка (scipy)',
)

add_table_to_doc(doc, table_scipy, 'Статистический ряд (scipy)')
add_table_to_doc(doc, chi_scipy_table, 'Критерий Пирсона (scipy)')
add_table_to_doc(doc, stats_scipy, 'Числовые характеристики (scipy)')

add_table_to_doc(doc, uniform_table, 'Критерий однородности')

doc.save('report_binom.docx')
